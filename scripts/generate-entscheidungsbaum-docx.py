#!/usr/bin/env python3
"""
Generates a Word document with the complete Abmahnchecker decision tree
for review purposes.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1C, 0x14, 0x08)

# ── Helper functions ──
def add_colored_text(paragraph, text, color_rgb, bold=False):
    run = paragraph.add_run(text)
    run.font.color.rgb = color_rgb
    run.bold = bold
    return run

def add_table_row(table, cells_data, bold_first=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9.5)
        if bold_first and i == 0:
            run.bold = True
        if bg_color:
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): bg_color,
            })
            shading.append(shd)
    return row

def set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shd)


# ══════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════

title = doc.add_heading('Abmahnchecker — Vollständiger Entscheidungsbaum', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('gekuendigt-abfindung.de  ·  /abmahnung-pruefen/')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x7A, 0x65, 0x28)
run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Dieses Dokument enthält den vollständigen Entscheidungsbaum des Abmahnchecker-Tools. ').font.size = Pt(10.5)
p.add_run('Es dient zur Prüfung und Kommentierung durch den Mandanten / Fachanwalt.').font.size = Pt(10.5)
doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 1. ÜBERSICHT
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('1. Übersicht', level=1)

doc.add_paragraph(
    'Der Abmahnchecker besteht aus 14 Fragen in 5 Blöcken (A–E). '
    'Einige Fragen werden nur unter bestimmten Bedingungen angezeigt (bedingte Logik). '
    'Jede Antwortoption hat ein Gewicht (Weight), das in die Gesamtbewertung einfließt, '
    'und kann optional Findings (Feststellungen mit BAG-Fundstellen) auslösen.'
)

doc.add_heading('Blöcke', level=2)
blocks = [
    ('A', 'Formelle Prüfung', '6 Fragen', 'Schriftform, Zugang, Zuständigkeit, Bestimmtheit, Hinweis- und Warnfunktion'),
    ('B', 'Materielle Prüfung', '4 Fragen', 'Sachverhalt, geschützte Rechte, Verschulden, AGG'),
    ('C', 'Sammelabmahnung', '2 Fragen', 'Mehrere Vorwürfe, unrichtiger Vorwurf in Sammelabmahnung'),
    ('D', 'Zeitliche Prüfung', '2 Fragen', 'Zeitablauf, Verwirkung'),
    ('E', 'Sonderfälle', '4 Fragen', 'Betriebsrat, Schwerbehinderung, Amtstätigkeit'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0]
for i, text in enumerate(['Block', 'Bezeichnung', 'Fragen', 'Themen']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

for b in blocks:
    add_table_row(table, b, bold_first=True)

doc.add_paragraph('')

doc.add_heading('Gewichtungen (Weights)', level=2)
weights = [
    ('KO', 'Knock-Out — 1× reicht für Verdict UNWIRKSAM', 'CC0000'),
    ('STRONG_NEGATIVE', 'Starkes negatives Indiz', 'E65100'),
    ('NEGATIVE', 'Negatives Indiz', 'F9A825'),
    ('NEUTRAL', 'Kein Einfluss auf Bewertung', '666666'),
    ('POSITIVE', 'Positives Indiz (für Wirksamkeit)', '2E7D32'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0]
for i, text in enumerate(['Weight', 'Bedeutung']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

for w_id, w_desc, w_color in weights:
    row = table.add_row()
    row.cells[0].text = ''
    p = row.cells[0].paragraphs[0]
    run = p.add_run(w_id)
    run.font.size = Pt(9.5)
    run.bold = True
    run.font.color.rgb = RGBColor(int(w_color[:2], 16), int(w_color[2:4], 16), int(w_color[4:6], 16))
    row.cells[1].text = w_desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 2. FRAGEN & ANTWORTOPTIONEN
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('2. Fragen & Antwortoptionen (vollständig)', level=1)

# Data structure mirroring questions.ts
questions = [
    # Block A
    {
        'id': 'A1', 'block': 'A', 'blockLabel': 'Formelle Prüfung',
        'text': 'Wurde die Abmahnung schriftlich erteilt?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'A1_ja', 'label': 'Ja, schriftlich', 'hint': 'Indiz: wirksam', 'hintTone': 'success', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A1_nein', 'label': 'Nein, mündlich', 'hint': 'Beweisproblem für Arbeitgeber', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': ['A1_muendlich']},
        ],
    },
    {
        'id': 'A2', 'block': 'A', 'blockLabel': 'Formelle Prüfung',
        'text': 'Ist die Abmahnung Ihnen tatsächlich zugegangen (z.\u00A0B. ausgehändigt, per Post, im Briefkasten)?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'A2_ja', 'label': 'Ja', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A2_nein', 'label': 'Nein, ich habe sie nicht erhalten', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['A2_kein_zugang']},
            {'id': 'A2_unsicher', 'label': 'Unsicher', 'hint': 'Beweislast Arbeitgeber', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': []},
        ],
    },
    {
        'id': 'A3', 'block': 'A', 'blockLabel': 'Formelle Prüfung',
        'text': 'Von wem wurde die Abmahnung ausgesprochen?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'A3_gf', 'label': 'Geschäftsführer / Vorstand / Personalleitung', 'hint': 'Indiz: wirksam', 'hintTone': 'success', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A3_vorgesetzter', 'label': 'Direkter Vorgesetzter mit Weisungsbefugnis', 'hint': 'Indiz: wirksam', 'hintTone': 'success', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A3_unbefugt', 'label': 'Eine Person ohne Weisungs- oder Personalbefugnis', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['A3_unzustaendig']},
            {'id': 'A3_unklar', 'label': 'Unklar', 'hint': 'Klärung empfohlen', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': []},
        ],
    },
    {
        'id': 'A4', 'block': 'A', 'blockLabel': 'Formelle Prüfung',
        'text': 'Beschreibt die Abmahnung das vorgeworfene Verhalten konkret nach Datum, Uhrzeit und Ort?',
        'help': 'Eine wirksame Abmahnung muss aus sich selbst heraus verständlich sein. Pauschale Vorwürfe wie \u201eStörung des Betriebsfriedens\u201c genügen nicht.',
        'showIf': None,
        'options': [
            {'id': 'A4_ja', 'label': 'Ja, alles konkret benannt', 'hint': 'Indiz: wirksam', 'hintTone': 'success', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A4_nein', 'label': 'Nein, nur pauschale Vorwürfe', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['A4_unbestimmt']},
            {'id': 'A4_teilweise', 'label': 'Teilweise', 'hint': 'Starkes Indiz: unwirksam', 'hintTone': 'warning', 'weight': 'STRONG_NEGATIVE', 'triggers': ['A4_teilweise']},
            {'id': 'A4_weiss_nicht', 'label': 'Weiß ich nicht', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
    {
        'id': 'A5', 'block': 'A', 'blockLabel': 'Formelle Prüfung',
        'text': 'Ist aus der Abmahnung erkennbar, welche konkrete arbeitsvertragliche Pflicht verletzt sein soll?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'A5_ja', 'label': 'Ja', 'hint': 'Indiz: wirksam', 'hintTone': 'success', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A5_nein', 'label': 'Nein', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['A5_keine_hinweisfunktion']},
            {'id': 'A5_unklar', 'label': 'Unklar', 'hint': 'Indiz: unwirksam', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': []},
        ],
    },
    {
        'id': 'A6', 'block': 'A', 'blockLabel': 'Formelle Prüfung',
        'text': 'Wird in der Abmahnung mit Konsequenzen (z.\u00A0B. Kündigung) für den Wiederholungsfall gedroht?',
        'help': 'Ohne Warnfunktion liegt nur eine Ermahnung vor \u2014 keine kündigungsrelevante Abmahnung.',
        'showIf': None,
        'options': [
            {'id': 'A6_ja', 'label': 'Ja, Kündigung oder arbeitsrechtliche Konsequenzen werden angedroht', 'hint': 'Indiz: wirksam', 'hintTone': 'success', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'A6_nein', 'label': 'Nein, nur Missbilligung ohne Drohung', 'hint': 'Nur Ermahnung, keine Abmahnung', 'hintTone': 'warning', 'weight': 'STRONG_NEGATIVE', 'triggers': ['A6_keine_warnfunktion']},
            {'id': 'A6_unklar', 'label': 'Unklar', 'hint': 'Klärung empfohlen', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': []},
        ],
    },
    # Block B
    {
        'id': 'B1', 'block': 'B', 'blockLabel': 'Materielle Prüfung',
        'text': 'Hat das in der Abmahnung beschriebene Verhalten tatsächlich so stattgefunden?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'B1_ja', 'label': 'Ja, vollständig', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'B1_teilweise', 'label': 'Teilweise \u2014 manche Punkte stimmen nicht', 'hint': 'Starkes Indiz: unwirksam', 'hintTone': 'warning', 'weight': 'STRONG_NEGATIVE', 'triggers': ['B1_teilweise']},
            {'id': 'B1_nein', 'label': 'Nein, der Sachverhalt ist falsch', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['B1_falsch']},
        ],
    },
    {
        'id': 'B2', 'block': 'B', 'blockLabel': 'Materielle Prüfung',
        'text': 'War das beanstandete Verhalten durch ein gesetzliches Recht gedeckt?',
        'help': 'Z.\u00A0B. ordnungsgemäße Krankmeldung, Ausübung von Betriebsratstätigkeit, Whistleblowing nach HinSchG, Streikteilnahme, Inanspruchnahme von Elternzeit/Urlaub, Beschwerde nach § 84 BetrVG.',
        'showIf': None,
        'options': [
            {'id': 'B2_krankmeldung', 'label': 'Ja, ordnungsgemäße Krankmeldung wird abgemahnt', 'hint': 'KO: § 612a BGB', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['B2_krankmeldung']},
            {'id': 'B2_br', 'label': 'Ja, Betriebsratstätigkeit / Amtsausübung', 'hint': 'KO: § 78 BetrVG', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['B2_br_taetigkeit']},
            {'id': 'B2_anderes', 'label': 'Ja, anderes geschütztes Recht (Elternzeit, Urlaub, Streik, Whistleblowing, Beschwerde)', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['B2_anderes_recht']},
            {'id': 'B2_nein', 'label': 'Nein', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
        ],
    },
    {
        'id': 'B3', 'block': 'B', 'blockLabel': 'Materielle Prüfung',
        'text': 'War das Verhalten Ihnen vorwerfbar (steuerbar und schuldhaft)?',
        'help': 'Krankheitsbedingte Minderleistung oder objektiv unmögliche Aufgaben sind nicht abmahnbar.',
        'showIf': None,
        'options': [
            {'id': 'B3_ja', 'label': 'Ja, ich hätte mich anders verhalten können', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'B3_nein', 'label': 'Nein, das Verhalten war nicht steuerbar', 'hint': 'KO-Kriterium', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['B3_kein_verschulden']},
            {'id': 'B3_unsicher', 'label': 'Unsicher', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
    {
        'id': 'B4', 'block': 'B', 'blockLabel': 'Materielle Prüfung',
        'text': 'Stellt die Abmahnung eine Diskriminierung im Sinne des AGG dar (Geschlecht, Alter, Religion, Behinderung, Herkunft, sex. Identität)?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'B4_nein', 'label': 'Nein', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'B4_ja', 'label': 'Ja', 'hint': 'KO: §§ 1, 7 AGG', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['B4_agg']},
            {'id': 'B4_unsicher', 'label': 'Unsicher', 'hint': 'Klärung empfohlen', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': []},
        ],
    },
    # Block C
    {
        'id': 'C1', 'block': 'C', 'blockLabel': 'Sammelabmahnung',
        'text': 'Werden in der Abmahnung MEHRERE unterschiedliche Pflichtverletzungen gerügt?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'C1_ja', 'label': 'Ja, mehrere Vorwürfe', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
            {'id': 'C1_nein', 'label': 'Nein, nur ein Vorwurf', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
        ],
    },
    {
        'id': 'C2', 'block': 'C', 'blockLabel': 'Sammelabmahnung',
        'text': 'Ist mindestens einer dieser Vorwürfe inhaltlich falsch oder nicht hinreichend konkretisiert?',
        'help': None,
        'showIf': 'Nur sichtbar, wenn C1 = "Ja, mehrere Vorwürfe"',
        'options': [
            {'id': 'C2_ja', 'label': 'Ja', 'hint': 'KO: Sammelabmahnung', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['C2_sammelabmahnung']},
            {'id': 'C2_nein', 'label': 'Nein, alle Vorwürfe treffen zu und sind konkret', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'C2_unsicher', 'label': 'Unsicher', 'hint': '', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': []},
        ],
    },
    # Block D
    {
        'id': 'D1', 'block': 'D', 'blockLabel': 'Zeitliche Prüfung',
        'text': 'Wie viel Zeit ist zwischen dem vorgeworfenen Verhalten und der Erteilung der Abmahnung vergangen?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'D1_4w', 'label': 'Bis 4 Wochen', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'D1_1_6', 'label': '1 bis 6 Monate', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
            {'id': 'D1_6_12', 'label': '6 bis 12 Monate', 'hint': 'Möglicher Verwirkungseinwand', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': ['D1_6_12_monate']},
            {'id': 'D1_12plus', 'label': 'Mehr als 12 Monate', 'hint': 'Starkes Indiz: Verwirkung', 'hintTone': 'warning', 'weight': 'STRONG_NEGATIVE', 'triggers': ['D1_ueber_12']},
            {'id': 'D1_weiss_nicht', 'label': 'Weiß ich nicht', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
    {
        'id': 'D2', 'block': 'D', 'blockLabel': 'Zeitliche Prüfung',
        'text': 'Hat der Arbeitgeber zwischen Vorfall und Abmahnung den Eindruck erweckt, die Sache sei erledigt (z.\u00A0B. Schweigen, Lob, Bonus, klärendes Gespräch ohne Abmahnungsvorbehalt)?',
        'help': None,
        'showIf': 'Nur sichtbar, wenn D1 = "6 bis 12 Monate" ODER "Mehr als 12 Monate"',
        'options': [
            {'id': 'D2_ja', 'label': 'Ja', 'hint': 'Verwirkungs-Konstellation', 'hintTone': 'warning', 'weight': 'STRONG_NEGATIVE', 'triggers': ['D2_verwirkung']},
            {'id': 'D2_nein', 'label': 'Nein', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
    # Block E
    {
        'id': 'E1', 'block': 'E', 'blockLabel': 'Sonderfälle',
        'text': 'Sind Sie Mitglied des Betriebsrats, Wahlvorstands, Sprecherausschusses oder der Schwerbehindertenvertretung?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'E1_ja', 'label': 'Ja', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
            {'id': 'E1_nein', 'label': 'Nein', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
    {
        'id': 'E2', 'block': 'E', 'blockLabel': 'Sonderfälle',
        'text': 'Bezieht sich der Abmahnungsvorwurf auf Ihre Amtstätigkeit (z.\u00A0B. Verhalten in der BR-Sitzung, Wahrnehmung von Aufgaben nach BetrVG)?',
        'help': None,
        'showIf': 'Nur sichtbar, wenn E1 = "Ja"',
        'options': [
            {'id': 'E2_ja', 'label': 'Ja, reine Amtstätigkeit', 'hint': 'KO: § 78 BetrVG', 'hintTone': 'danger', 'weight': 'KO', 'triggers': ['E2_amtstaetigkeit']},
            {'id': 'E2_nein', 'label': 'Nein, normale arbeitsvertragliche Pflicht', 'hint': 'ok', 'hintTone': 'neutral', 'weight': 'POSITIVE', 'triggers': []},
            {'id': 'E2_mischfall', 'label': 'Mischfall', 'hint': 'Starkes Indiz: unwirksam', 'hintTone': 'warning', 'weight': 'STRONG_NEGATIVE', 'triggers': ['E2_mischfall']},
        ],
    },
    {
        'id': 'E3', 'block': 'E', 'blockLabel': 'Sonderfälle',
        'text': 'Sind Sie schwerbehindert oder gleichgestellt?',
        'help': None,
        'showIf': None,
        'options': [
            {'id': 'E3_ja', 'label': 'Ja', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
            {'id': 'E3_nein', 'label': 'Nein', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
    {
        'id': 'E4', 'block': 'E', 'blockLabel': 'Sonderfälle',
        'text': 'Hat der Vorwurf einen Bezug zu Ihrer Behinderung?',
        'help': None,
        'showIf': 'Nur sichtbar, wenn E3 = "Ja"',
        'options': [
            {'id': 'E4_ja', 'label': 'Ja', 'hint': 'SBV-Anhörung prüfen', 'hintTone': 'warning', 'weight': 'NEGATIVE', 'triggers': ['E4_sbv_anhoerung']},
            {'id': 'E4_nein', 'label': 'Nein', 'hint': '', 'hintTone': 'neutral', 'weight': 'NEUTRAL', 'triggers': []},
        ],
    },
]

# Weight → color mapping for table
weight_colors = {
    'KO': 'FFCCCC',
    'STRONG_NEGATIVE': 'FFE0CC',
    'NEGATIVE': 'FFF3CC',
    'NEUTRAL': 'F0F0F0',
    'POSITIVE': 'D4EDDA',
}

current_block = None
for q in questions:
    # New block heading
    if q['block'] != current_block:
        current_block = q['block']
        doc.add_heading(f"Block {current_block} — {q['blockLabel']}", level=2)

    # Question heading
    doc.add_heading(f"{q['id']}: {q['text']}", level=3)

    # Help text
    if q.get('help'):
        p = doc.add_paragraph()
        run = p.add_run('Erklärungstext: ')
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run = p.add_run(q['help'])
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True

    # Conditional logic
    if q.get('showIf'):
        p = doc.add_paragraph()
        run = p.add_run('\u26A0 Bedingte Anzeige: ')
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
        run = p.add_run(q['showIf'])
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

    # Options table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, text in enumerate(['Option-ID', 'Antworttext', 'Hinweis (Tooltip)', 'Weight', 'Löst Findings aus']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
        set_cell_bg(hdr.cells[i], 'F0EAD9')

    for opt in q['options']:
        row = table.add_row()
        row.cells[0].text = opt['id']
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)

        row.cells[1].text = opt['label']
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(8.5)

        row.cells[2].text = opt['hint'] or '—'
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(8.5)

        # Weight with color
        row.cells[3].text = ''
        p = row.cells[3].paragraphs[0]
        run = p.add_run(opt['weight'])
        run.font.size = Pt(8.5)
        run.bold = True
        bg = weight_colors.get(opt['weight'], 'FFFFFF')
        set_cell_bg(row.cells[3], bg)

        triggers_text = ', '.join(opt['triggers']) if opt['triggers'] else '—'
        row.cells[4].text = triggers_text
        row.cells[4].paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph('')  # spacing

    # Add comment area
    p = doc.add_paragraph()
    run = p.add_run(f'Kommentar zu {q["id"]}: ')
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run('_______________________________________________')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 3. FINDINGS (Feststellungen)
# ══════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading('3. Findings — Feststellungen mit BAG-Fundstellen', level=1)

doc.add_paragraph(
    'Jede Finding wird durch einen Trigger aus einer Antwortoption ausgelöst. '
    'Severity "KO" bedeutet: diese einzelne Feststellung allein führt zum Verdict UNWIRKSAM. '
    'Severity "INDIZ" fließt in die Gesamtbewertung ein.'
)

findings = [
    {'id': 'A1_muendlich', 'severity': 'INDIZ', 'title': 'Mündliche Abmahnung — schwere Beweisprobleme für den Arbeitgeber', 'body': 'Eine mündliche Abmahnung ist zwar grundsätzlich wirksam (§ 126 BGB findet keine Anwendung), praktisch jedoch im Streitfall kaum beweisbar. Der Arbeitgeber trägt die volle Beweislast für Inhalt und Zugang.', 'citation': 'BAG, Urt. v. 19.02.2009 – 2 AZR 603/07'},
    {'id': 'A2_kein_zugang', 'severity': 'KO', 'title': 'Abmahnung nicht zugegangen', 'body': 'Die Abmahnung ist eine empfangsbedürftige Willenserklärung. Ohne Zugang nach § 130 BGB entfaltet sie keinerlei Wirkung.', 'citation': '§ 130 BGB'},
    {'id': 'A3_unzustaendig', 'severity': 'KO', 'title': 'Unzuständige Person hat abgemahnt', 'body': 'Abmahnberechtigt sind nur Kündigungsberechtigte und Vorgesetzte mit Direktionsrecht hinsichtlich der Arbeitsleistung. Eine Abmahnung von einer nicht weisungsberechtigten Person ist unwirksam.', 'citation': 'BAG AP Nr. 3 zu § 1 KSchG 1969 Verhaltensbedingte Kündigung'},
    {'id': 'A4_unbestimmt', 'severity': 'KO', 'title': 'Fehlende Bestimmtheit des Vorwurfs', 'body': 'Die Abmahnung enthält nur pauschale Vorwürfe. Eine wirksame Abmahnung muss das konkrete Verhalten nach Datum, Uhrzeit und Ort bezeichnen, sodass sie aus sich selbst heraus verständlich ist.', 'citation': 'BAG, Urt. v. 27.11.2008 – 2 AZR 675/07'},
    {'id': 'A4_teilweise', 'severity': 'INDIZ', 'title': 'Teilweise unbestimmte Vorwürfe', 'body': 'Soweit Vorwürfe pauschal bleiben, sind sie unwirksam. Zweifel gehen zu Lasten des Arbeitgebers.', 'citation': 'BAG, Urt. v. 27.11.2008 – 2 AZR 675/07'},
    {'id': 'A5_keine_hinweisfunktion', 'severity': 'KO', 'title': 'Hinweisfunktion fehlt', 'body': 'Es ist nicht erkennbar, welche konkrete arbeitsvertragliche Pflicht verletzt worden sein soll. Die Hinweisfunktion verlangt, dass dem Arbeitnehmer der Pflichtverletzungs-Charakter seines Verhaltens deutlich gemacht wird.', 'citation': 'BAG, Urt. v. 19.04.2012 – 2 AZR 258/11'},
    {'id': 'A6_keine_warnfunktion', 'severity': 'KO', 'title': 'Warnfunktion fehlt — keine Abmahnung im Rechtssinne', 'body': 'Ohne Androhung arbeitsrechtlicher Konsequenzen für den Wiederholungsfall liegt nur eine Ermahnung/Verwarnung vor. Diese kann keine Vorstufe einer verhaltensbedingten Kündigung sein.', 'citation': 'BAG, Urt. v. 19.04.2012 – 2 AZR 258/11'},
    {'id': 'B1_falsch', 'severity': 'KO', 'title': 'Unrichtige Tatsachenbehauptung', 'body': 'Eine Abmahnung mit unzutreffendem Sachverhalt ist insgesamt aus der Personalakte zu entfernen, sofern der Arbeitgeber den Vorwurf nicht beweisen kann.', 'citation': 'BAG, Urt. v. 12.08.2010 – 2 AZR 593/09'},
    {'id': 'B1_teilweise', 'severity': 'INDIZ', 'title': 'Teilweise unrichtige Tatsachendarstellung', 'body': 'Soweit einzelne Vorwürfe nicht zutreffen, sind sie aus der Abmahnung zu entfernen. Bei Sammelabmahnungen führt dies zur Gesamtunwirksamkeit.', 'citation': 'BAG, Urt. v. 13.03.1991 – 5 AZR 133/90'},
    {'id': 'B2_krankmeldung', 'severity': 'KO', 'title': 'Verstoß gegen das Maßregelungsverbot — Krankmeldung', 'body': 'Die ordnungsgemäße Krankmeldung ist Ausübung eines gesetzlichen Rechts (EFZG). Eine Abmahnung wegen der Krankheit selbst verstößt gegen § 612a BGB.', 'citation': 'LAG Hessen, Urt. v. 28.03.2025 – 10 SLa 916/24; § 612a BGB'},
    {'id': 'B2_br_taetigkeit', 'severity': 'KO', 'title': 'Abmahnung wegen Betriebsratstätigkeit unzulässig', 'body': 'Eine individualrechtliche Abmahnung wegen reiner betriebsverfassungsrechtlicher Amtspflichtverletzungen ist nach ständiger BAG-Rechtsprechung ausgeschlossen. Sanktion ist allenfalls § 23 Abs. 1 BetrVG.', 'citation': 'BAG, Beschl. v. 09.09.2015 – 7 ABR 69/13'},
    {'id': 'B2_anderes_recht', 'severity': 'KO', 'title': 'Verstoß gegen das Maßregelungsverbot', 'body': 'Die Abmahnung sanktioniert die Ausübung eines gesetzlich geschützten Rechts und verstößt gegen § 612a BGB bzw. spezielle Maßregelungsverbote (§ 5 TzBfG, § 16 AGG, § 78 S. 2 BetrVG, § 36 HinSchG).', 'citation': '§ 612a BGB'},
    {'id': 'B3_kein_verschulden', 'severity': 'KO', 'title': 'Fehlendes Verschulden', 'body': 'Abmahnungen setzen ein steuerbares und schuldhaftes Verhalten voraus. Krankheitsbedingte Minderleistungen oder objektiv unmögliche Aufgaben sind nicht abmahnbar.', 'citation': 'BAG (st. Rspr.)'},
    {'id': 'B4_agg', 'severity': 'KO', 'title': 'AGG-Verstoß', 'body': 'Die Abmahnung knüpft an ein Diskriminierungsmerkmal nach § 1 AGG an und verstößt gegen §§ 7, 16 AGG.', 'citation': '§§ 1, 7, 16 AGG'},
    {'id': 'C2_sammelabmahnung', 'severity': 'KO', 'title': 'Sammelabmahnung mit unrichtigem Vorwurf', 'body': 'Die Abmahnung enthält mehrere Vorwürfe, von denen mindestens einer inhaltlich nicht zutrifft oder unkonkret bleibt. Nach gefestigter BAG-Rechtsprechung ist die gesamte Abmahnung in einem solchen Fall aus der Personalakte zu entfernen.', 'citation': 'BAG, Urt. v. 13.03.1991 – 5 AZR 133/90'},
    {'id': 'D1_6_12_monate', 'severity': 'INDIZ', 'title': 'Verzögerte Erteilung — möglicher Verwirkungseinwand', 'body': 'Zwischen Vorfall und Abmahnung liegen mehr als 6 Monate. In Verbindung mit einem Umstandsmoment kann dies zur Verwirkung nach § 242 BGB führen.', 'citation': 'LAG Nürnberg, Urt. v. 14.06.2005 – 6 Sa 367/05; ArbG Mainz, Urt. v. 19.01.2023 – 7 Ca 581/22'},
    {'id': 'D1_ueber_12', 'severity': 'INDIZ', 'title': 'Erhebliche Verzögerung — starkes Verwirkungsindiz', 'body': 'Eine Abmahnung mehr als 12 Monate nach dem Vorfall ist regelmäßig verwirkt, sofern kein ausdrücklicher Vorbehalt erklärt wurde.', 'citation': '§ 242 BGB'},
    {'id': 'D2_verwirkung', 'severity': 'KO', 'title': 'Verwirkung nach § 242 BGB', 'body': 'Der Arbeitgeber hat durch sein Verhalten den Eindruck erweckt, der Vorgang sei erledigt. Bei einem Zeitablauf von mehr als 6 Monaten und hinzutretendem Umstandsmoment ist das Abmahnungsrecht verwirkt.', 'citation': '§ 242 BGB; LAG Nürnberg 6 Sa 367/05'},
    {'id': 'E2_amtstaetigkeit', 'severity': 'KO', 'title': 'Abmahnung wegen Amtstätigkeit unzulässig', 'body': 'Reine betriebsverfassungsrechtliche Amtspflichtverletzungen rechtfertigen keine individualrechtliche Abmahnung. Sanktion ist allein § 23 Abs. 1 BetrVG.', 'citation': 'BAG, Beschl. v. 09.09.2015 – 7 ABR 69/13'},
    {'id': 'E2_mischfall', 'severity': 'INDIZ', 'title': 'Mischfall Amtstätigkeit / Arbeitspflicht', 'body': 'Eine Abmahnung wegen Amtspflichtverletzung mit Androhung individualrechtlicher Sanktionen ist unwirksam, sofern nicht klar getrennt wird.', 'citation': 'BAG, Beschl. v. 09.09.2015 – 7 ABR 69/13'},
    {'id': 'E4_sbv_anhoerung', 'severity': 'INDIZ', 'title': 'Schwerbehindertenvertretung möglicherweise nicht angehört', 'body': 'Bei einem Bezug des Vorwurfs zur Behinderung ist die SBV gemäß § 178 Abs. 2 SGB IX vor Erteilung anzuhören (h.\u00A0M.). Bei unterlassener Anhörung droht formelle Unwirksamkeit.', 'citation': '§ 178 Abs. 2 SGB IX'},
]

for f in findings:
    severity_label = f['severity']
    severity_color = 'FFCCCC' if severity_label == 'KO' else 'FFF3CC'

    # Finding title with severity badge
    p = doc.add_paragraph()
    run = p.add_run(f"[{severity_label}] ")
    run.bold = True
    run.font.size = Pt(10)
    if severity_label == 'KO':
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    else:
        run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

    run = p.add_run(f['title'])
    run.bold = True
    run.font.size = Pt(10)

    # Trigger ID
    p = doc.add_paragraph()
    run = p.add_run(f"Trigger-ID: {f['id']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Body
    p = doc.add_paragraph(f['body'])
    p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.size = Pt(9.5)

    # Citation
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('Fundstelle: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(f['citation'])
    run.font.size = Pt(9)
    run.italic = True

    # Comment line
    p = doc.add_paragraph()
    run = p.add_run(f'Kommentar: ')
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run('_______________________________________________')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 4. BEWERTUNGSLOGIK (Evaluator)
# ══════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading('4. Bewertungslogik — Wie wird das Verdict ermittelt?', level=1)

doc.add_paragraph(
    'Die Bewertungslogik sammelt alle Gewichtungen (Weights) und Triggers aus den '
    'beantworteten Fragen und ermittelt daraus das Gesamtverdikt.'
)

doc.add_heading('Schritt 1: Zählung', level=2)
doc.add_paragraph(
    'Für jede beantwortete Frage wird das Weight der gewählten Option gezählt:\n'
    '  \u2022 ko: Anzahl KO-Optionen\n'
    '  \u2022 strongNegative: Anzahl STRONG_NEGATIVE-Optionen\n'
    '  \u2022 negative: Anzahl NEGATIVE-Optionen\n'
    '  \u2022 positive: Anzahl POSITIVE-Optionen\n'
    '  \u2022 NEUTRAL wird nicht gezählt.'
)

doc.add_heading('Schritt 2: Findings sammeln', level=2)
doc.add_paragraph(
    'Alle Trigger-IDs der gewählten Optionen werden gesammelt. '
    'Für jeden Trigger wird das zugehörige Finding nachgeschlagen. '
    'Doppelte werden ignoriert. Findings werden sortiert: KO zuerst, dann INDIZ.'
)

doc.add_heading('Schritt 3: Verdict bestimmen', level=2)

# Verdict logic as table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0]
for i, text in enumerate(['Priorität', 'Bedingung', 'Verdict']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

verdict_rules = [
    ('1 (höchste)', 'KO-Findings >= 1', 'UNWIRKSAM'),
    ('2', 'STRONG_NEGATIVE >= 2', 'UNWIRKSAM'),
    ('3', 'STRONG_NEGATIVE >= 1 UND NEGATIVE >= 2', 'UNWIRKSAM'),
    ('4', 'STRONG_NEGATIVE >= 1 ODER NEGATIVE >= 2', 'UNSICHER'),
    ('5 (niedrigste)', 'Keiner der oberen trifft zu', 'WIRKSAM'),
]

verdict_colors = {
    'UNWIRKSAM': 'FFCCCC',
    'UNSICHER': 'FFF3CC',
    'WIRKSAM': 'D4EDDA',
}

for rule in verdict_rules:
    row = table.add_row()
    for i, text in enumerate(rule):
        row.cells[i].text = text
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    if rule[2] in verdict_colors:
        set_cell_bg(row.cells[2], verdict_colors[rule[2]])

doc.add_paragraph('')

doc.add_heading('Verdict-Darstellung auf der Ergebnisseite', level=2)

verdicts_display = [
    ('UNWIRKSAM', 'Rot', 'Ihre Abmahnung weist erhebliche Mängel auf. Sie haben gute Chancen, die Entfernung aus der Personalakte zu verlangen.'),
    ('UNSICHER', 'Amber/Gelb', 'Ihre Abmahnung weist angreifbare Punkte auf. Eine anwaltliche Prüfung ist empfehlenswert.'),
    ('WIRKSAM', 'Grün/Teal', 'Ihre Abmahnung erscheint nach den vorliegenden Angaben formell und inhaltlich wirksam. Eine anwaltliche Gegendarstellung kann dennoch sinnvoll sein.'),
]

for v_id, v_color, v_text in verdicts_display:
    p = doc.add_paragraph()
    run = p.add_run(f'{v_id} ({v_color}): ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(v_text)
    run.font.size = Pt(10)

doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 5. BEDINGTE LOGIK — ZUSAMMENFASSUNG
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('5. Bedingte Logik — Zusammenfassung', level=1)

doc.add_paragraph(
    'Folgende Fragen werden nur unter bestimmten Bedingungen angezeigt:'
)

cond_table = doc.add_table(rows=1, cols=3)
cond_table.style = 'Table Grid'
hdr = cond_table.rows[0]
for i, text in enumerate(['Frage', 'Wird angezeigt wenn...', 'Sonst']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

conditionals = [
    ('C2', 'C1 = "Ja, mehrere Vorwürfe"', 'Wird übersprungen'),
    ('D2', 'D1 = "6 bis 12 Monate" ODER "Mehr als 12 Monate"', 'Wird übersprungen'),
    ('E2', 'E1 = "Ja"', 'Wird übersprungen'),
    ('E4', 'E3 = "Ja"', 'Wird übersprungen'),
]

for c in conditionals:
    add_table_row(cond_table, c, bold_first=True)

doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 6. TRIGGER→FINDING ZUORDNUNG
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('6. Trigger → Finding Zuordnung (Querverweisliste)', level=1)

doc.add_paragraph(
    'Diese Tabelle zeigt, welche Antwortoption welches Finding auslöst:'
)

ref_table = doc.add_table(rows=1, cols=5)
ref_table.style = 'Table Grid'
hdr = ref_table.rows[0]
for i, text in enumerate(['Frage', 'Option', 'Trigger-ID', 'Finding-Severity', 'Finding-Titel']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

# Build finding lookup
findings_lookup = {f['id']: f for f in findings}

for q in questions:
    for opt in q['options']:
        if opt['triggers']:
            for trigger in opt['triggers']:
                f = findings_lookup.get(trigger, {})
                row = ref_table.add_row()
                cells = [
                    q['id'],
                    opt['label'][:40] + ('...' if len(opt['label']) > 40 else ''),
                    trigger,
                    f.get('severity', '?'),
                    f.get('title', '?')[:50] + ('...' if len(f.get('title', '')) > 50 else ''),
                ]
                for i, text in enumerate(cells):
                    row.cells[i].text = text
                    row.cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                if f.get('severity') == 'KO':
                    set_cell_bg(row.cells[3], 'FFCCCC')

doc.add_paragraph('')
doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Ende des Entscheidungsbaums —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('APOS Legal  ·  gekuendigt-abfindung.de  ·  Generiert am: April 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_dir = os.path.expanduser('~/Desktop')
output_path = os.path.join(output_dir, 'Abmahnchecker-Entscheidungsbaum.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
