#!/usr/bin/env python3
"""
Generates a Word document with the complete Aufhebungsvertrag-Checker decision tree.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1C, 0x14, 0x08)

def set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shd)

def add_table_row(table, cells_data, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9.5)
        if bold_first and i == 0:
            run.bold = True
    return row

# ══════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════

title = doc.add_heading('Aufhebungsvertrag-Checker — Vollständiger Entscheidungsbaum', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('gekuendigt-abfindung.de  ·  /aufhebungsvertrag-pruefen/')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x7A, 0x65, 0x28)
run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Dieses Dokument enthält den vollständigen Entscheidungsbaum des Aufhebungsvertrag-Checkers. ').font.size = Pt(10.5)
p.add_run('Es dient zur Prüfung und Kommentierung durch den Mandanten / Fachanwalt.').font.size = Pt(10.5)
doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 1. ÜBERSICHT
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('1. Übersicht', level=1)

doc.add_paragraph(
    'Der Aufhebungsvertrag-Checker besteht aus 19 Fragen in 8 Kategorien. '
    'Einige Fragen werden nur unter bestimmten Bedingungen angezeigt (bedingte Logik). '
    'Jede Antwortoption erhält eine Ampelfarbe (grün/gelb/rot), die in die Gesamtbewertung einfließt. '
    'Bestimmte Antworten lösen eine InfoBox mit rechtlichem Hinweis aus.'
)

doc.add_heading('Kategorien', level=2)
categories = [
    ('1', 'Abfindung', 'S1, S2, S3, S19', 'Abfindung vorhanden? Höhe? Zeitpunkt? Turboklausel?'),
    ('2', 'Vergütung & Urlaub', 'S4, S5', 'Bonus/Provision geregelt? Resturlaub geregelt?'),
    ('3', 'Fristen & Sperrzeit', 'S6, S7, S8', 'Widerrufsfrist? Sperrzeit-Hinweis? Ausschlussfristen?'),
    ('4', 'Freistellung', 'S9, S10', 'Art der Freistellung? PKV-Thema?'),
    ('5', 'Zeugnis', 'S11, S12', 'Zeugnis vereinbart? Ausstellungsdatum?'),
    ('6', 'Wettbewerbsverbot', 'S13, S14', 'Wettbewerbsverbot? Karenzentschädigung?'),
    ('7', 'Besonderer Schutz', 'S15, S16', 'Sonderschutz? Stellen einbezogen?'),
    ('8', 'Verhandlung & Druck', 'S17, S18', 'Bedenkzeit? Druck/Drohung?'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0]
for i, text in enumerate(['Kat.', 'Bezeichnung', 'Fragen', 'Themen']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

for c in categories:
    add_table_row(table, c, bold_first=True)

doc.add_paragraph('')

doc.add_heading('Ampel-Bewertungssystem', level=2)

ampel_items = [
    ('Grün (2 Punkte)', 'Günstige / faire Klausel', 'D4EDDA'),
    ('Gelb (1 Punkt)', 'Neutral / unsicher / verbesserungswürdig', 'FFF3CC'),
    ('Rot (0 Punkte)', 'Problematisch / fehlend / nachteilig', 'FFCCCC'),
    ('Skip (nicht gezählt)', 'Frage ist nicht relevant (wird übersprungen)', 'F0F0F0'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0]
for i, text in enumerate(['Ampelfarbe', 'Bedeutung']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

for label, desc, color in ampel_items:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    row.cells[0].paragraphs[0].runs[0].bold = True
    set_cell_bg(row.cells[0], color)
    row.cells[1].text = desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 2. FRAGEN & ANTWORTOPTIONEN
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('2. Fragen & Antwortoptionen (vollständig)', level=1)

# Color map
color_map = {
    'green': 'D4EDDA',
    'yellow': 'FFF3CC',
    'red': 'FFCCCC',
    'skip': 'F0F0F0',
}

color_labels = {
    'green': 'GRÜN',
    'yellow': 'GELB',
    'red': 'ROT',
    'skip': 'SKIP',
}

# Full question data
questions = [
    # Kategorie 1: Abfindung
    {
        'id': 'S1', 'category': 1, 'catLabel': 'Abfindung',
        'text': 'Enthält der Aufhebungsvertrag eine Abfindung?',
        'showIf': None,
        'options': [
            {'label': 'Ja, eine Abfindung ist vereinbart', 'key': 'abfindungVorhanden', 'value': 'ja', 'color': 'green', 'next': 'S2', 'infoBox': None},
            {'label': 'Nein, keine Abfindung vorgesehen', 'key': 'abfindungVorhanden', 'value': 'nein', 'color': 'red', 'next': 'S19', 'infoBox': 'Kein gesetzlicher Anspruch auf Abfindung, aber bei Aufhebungsverträgen verhandelbar. Faustformel: 0,5 Monatsgehälter × Beschäftigungsjahre.'},
            {'label': 'Ich bin unsicher', 'key': 'abfindungVorhanden', 'value': 'unsicher', 'color': 'yellow', 'next': 'S19', 'infoBox': None},
        ],
    },
    {
        'id': 'S2', 'category': 1, 'catLabel': 'Abfindung',
        'text': 'Wie hoch ist die Abfindung im Verhältnis zu Gehalt und Betriebszugehörigkeit?',
        'showIf': 'Nur sichtbar, wenn S1 = "Ja, eine Abfindung ist vereinbart"',
        'options': [
            {'label': 'Über 0,5 Monatsgehälter pro Beschäftigungsjahr', 'key': 'abfindungHoehe', 'value': 'ueber', 'color': 'green', 'next': 'S3', 'infoBox': None},
            {'label': 'Genau 0,5 Monatsgehälter pro Beschäftigungsjahr', 'key': 'abfindungHoehe', 'value': 'genau', 'color': 'green', 'next': 'S3', 'infoBox': None},
            {'label': 'Unter 0,5 Monatsgehälter pro Beschäftigungsjahr', 'key': 'abfindungHoehe', 'value': 'unter', 'color': 'red', 'next': 'S3', 'infoBox': None},
            {'label': 'Ich weiß es nicht', 'key': 'abfindungHoehe', 'value': 'unsicher', 'color': 'yellow', 'next': 'S3', 'infoBox': None},
        ],
    },
    {
        'id': 'S3', 'category': 1, 'catLabel': 'Abfindung',
        'text': 'Wann wird die Abfindung ausgezahlt?',
        'showIf': 'Nur sichtbar, wenn S1 = "Ja"',
        'options': [
            {'label': 'Mit letztem Gehalt / zum Beendigungsdatum', 'key': 'abfindungZeitpunkt', 'value': 'sofort', 'color': 'green', 'next': 'S19', 'infoBox': None},
            {'label': 'Innerhalb von 4 Wochen nach Beendigung', 'key': 'abfindungZeitpunkt', 'value': '4wochen', 'color': 'green', 'next': 'S19', 'infoBox': None},
            {'label': 'Erst nach mehreren Monaten', 'key': 'abfindungZeitpunkt', 'value': 'spaet', 'color': 'yellow', 'next': 'S19', 'infoBox': None},
            {'label': 'Keine Regelung zum Auszahlungszeitpunkt', 'key': 'abfindungZeitpunkt', 'value': 'keine', 'color': 'red', 'next': 'S19', 'infoBox': None},
        ],
    },
    {
        'id': 'S19', 'category': 1, 'catLabel': 'Abfindung',
        'text': 'Enthält der Vertrag eine Turboklausel (Sprinterklausel)?',
        'showIf': None,
        'help': 'Eine Turboklausel gibt Ihnen das Recht, das Arbeitsverhältnis vor dem vereinbarten Beendigungsdatum zu beenden — die ersparte Vergütung wird dann als zusätzliche Abfindung ausgezahlt.',
        'options': [
            {'label': 'Ja, eine Turboklausel ist vereinbart', 'key': 'turboklausel', 'value': 'ja', 'color': 'green', 'next': 'S4', 'infoBox': None},
            {'label': 'Nein, keine Turboklausel enthalten', 'key': 'turboklausel', 'value': 'nein', 'color': 'yellow', 'next': 'S4', 'infoBox': 'Eine Turboklausel kann für Sie vorteilhaft sein: Finden Sie vor dem Beendigungsdatum eine neue Stelle, können Sie früher wechseln und erhalten die restliche Vergütung als zusätzliche Zahlung. Es lohnt sich, diese Klausel nachzuverhandeln.'},
            {'label': 'Ich weiß es nicht', 'key': 'turboklausel', 'value': 'unsicher', 'color': 'yellow', 'next': 'S4', 'infoBox': None},
        ],
    },
    # Kategorie 2: Vergütung & Urlaub
    {
        'id': 'S4', 'category': 2, 'catLabel': 'Vergütung & Urlaub',
        'text': 'Sind ausstehende Vergütungsbestandteile geregelt? (Bonus, Provision, Tantieme, Überstunden)',
        'showIf': None,
        'options': [
            {'label': 'Ja, alles explizit geregelt', 'key': 'bonusGeregelt', 'value': 'ja', 'color': 'green', 'next': 'S5', 'infoBox': None},
            {'label': 'Teilweise geregelt', 'key': 'bonusGeregelt', 'value': 'teilweise', 'color': 'yellow', 'next': 'S5', 'infoBox': None},
            {'label': 'Nein, keine Regelung', 'key': 'bonusGeregelt', 'value': 'nein', 'color': 'red', 'next': 'S5', 'infoBox': 'Ohne ausdrückliche Regelung können Bonus- und Provisionsansprüche nach Vertragsende verfallen. Bestehen Sie auf einer klaren Klausel.'},
            {'label': 'Ich habe keine solchen Ansprüche', 'key': 'bonusGeregelt', 'value': 'keine_ansprueche', 'color': 'green', 'next': 'S5', 'infoBox': None},
        ],
    },
    {
        'id': 'S5', 'category': 2, 'catLabel': 'Vergütung & Urlaub',
        'text': 'Ist Ihr Resturlaub im Vertrag geregelt?',
        'showIf': None,
        'options': [
            {'label': 'Ja — Freistellung unter Urlaubsanrechnung', 'key': 'urlaubGeregelt', 'value': 'freistellung', 'color': 'green', 'next': 'S6', 'infoBox': None},
            {'label': 'Ja — Urlaubsabgeltung in Geld vereinbart', 'key': 'urlaubGeregelt', 'value': 'abgeltung', 'color': 'green', 'next': 'S6', 'infoBox': None},
            {'label': 'Nein, Resturlaub nicht geregelt', 'key': 'urlaubGeregelt', 'value': 'nein', 'color': 'red', 'next': 'S6', 'infoBox': 'Resturlaub muss entweder gewährt oder finanziell abgegolten werden (§7 Abs. 4 BUrlG). Ohne Regelung verlieren Sie möglicherweise Ihren Anspruch.'},
        ],
    },
    # Kategorie 3: Fristen & Sperrzeit
    {
        'id': 'S6', 'category': 3, 'catLabel': 'Fristen & Sperrzeit',
        'text': 'Enthält der Vertrag eine Widerrufsfrist?',
        'showIf': None,
        'options': [
            {'label': 'Ja, mindestens 1 Woche Widerrufsfrist', 'key': 'widerrufsfrist', 'value': 'ja_1woche', 'color': 'green', 'next': 'S7', 'infoBox': None},
            {'label': 'Ja, weniger als 1 Woche', 'key': 'widerrufsfrist', 'value': 'ja_kurz', 'color': 'yellow', 'next': 'S7', 'infoBox': None},
            {'label': 'Nein, keine Widerrufsfrist enthalten', 'key': 'widerrufsfrist', 'value': 'nein', 'color': 'red', 'next': 'S7', 'infoBox': 'Eine Widerrufsfrist ist gesetzlich nicht vorgeschrieben, aber ein wichtiges Qualitätsmerkmal. Das BAG (6 AZR 75/18) hat das Gebot fairen Verhandelns etabliert. Ohne jede Bedenkzeit kann der Vertrag unter Umständen anfechtbar sein.'},
        ],
    },
    {
        'id': 'S7', 'category': 3, 'catLabel': 'Fristen & Sperrzeit',
        'text': 'Wurden Sie auf die mögliche Sperrzeit beim Arbeitslosengeld hingewiesen?',
        'showIf': None,
        'options': [
            {'label': 'Ja, ich wurde informiert und habe dem zugestimmt', 'key': 'sperrzeitHinweis', 'value': 'ja', 'color': 'green', 'next': 'S8', 'infoBox': None},
            {'label': 'Nein, kein Hinweis im Vertrag enthalten', 'key': 'sperrzeitHinweis', 'value': 'nein', 'color': 'red', 'next': 'S8', 'infoBox': 'Achtung: Bei Aufhebungsverträgen droht regelmäßig eine 12-wöchige Sperrzeit beim ALG I (§159 SGB III). Eine Abfindung vermeidet diese nicht automatisch. Lassen Sie prüfen, ob eine Formulierung als „betriebsbedingte" Beendigung möglich ist.'},
            {'label': 'Ich weiß es nicht', 'key': 'sperrzeitHinweis', 'value': 'unsicher', 'color': 'yellow', 'next': 'S8', 'infoBox': None},
        ],
    },
    {
        'id': 'S8', 'category': 3, 'catLabel': 'Fristen & Sperrzeit',
        'text': 'Enthält der Vertrag Ausschlussfristen für weitere Ansprüche?',
        'showIf': None,
        'options': [
            {'label': 'Ja, mit klarer Frist (mind. 3 Monate)', 'key': 'ausschlussfrist', 'value': 'ja_3m', 'color': 'green', 'next': 'S9', 'infoBox': None},
            {'label': 'Ja, aber Frist unter 3 Monate', 'key': 'ausschlussfrist', 'value': 'ja_kurz', 'color': 'red', 'next': 'S9', 'infoBox': 'Sehr kurze Ausschlussfristen können Sie unter Druck setzen, Ansprüche schnell geltend zu machen. Prüfen Sie welche Ansprüche davon betroffen sind.'},
            {'label': 'Keine Ausschlussfrist enthalten', 'key': 'ausschlussfrist', 'value': 'keine', 'color': 'yellow', 'next': 'S9', 'infoBox': None},
            {'label': 'Ich weiß es nicht', 'key': 'ausschlussfrist', 'value': 'unsicher', 'color': 'yellow', 'next': 'S9', 'infoBox': None},
        ],
    },
    # Kategorie 4: Freistellung
    {
        'id': 'S9', 'category': 4, 'catLabel': 'Freistellung',
        'text': 'Wie ist Ihre Freistellung bis Vertragsende geregelt?',
        'showIf': None,
        'options': [
            {'label': 'Bezahlte Freistellung unter voller Vergütungsfortzahlung', 'key': 'freistellungArt', 'value': 'bezahlt', 'color': 'green', 'next': 'S10', 'infoBox': None},
            {'label': 'Ich muss bis Vertragsende weiterarbeiten', 'key': 'freistellungArt', 'value': 'weiterarbeiten', 'color': 'yellow', 'next': 'S10', 'infoBox': None},
            {'label': 'Unbezahlte Freistellung vereinbart', 'key': 'freistellungArt', 'value': 'unbezahlt', 'color': 'red', 'next': 'S11', 'infoBox': 'Bei unbezahlter Freistellung entfällt die Pflichtversicherung in der GKV. Sie müssen sich freiwillig versichern — prüfen Sie Kosten und Fristen unverzüglich.'},
            {'label': 'Keine Regelung zur Freistellung', 'key': 'freistellungArt', 'value': 'keine', 'color': 'red', 'next': 'S11', 'infoBox': None},
        ],
    },
    {
        'id': 'S10', 'category': 4, 'catLabel': 'Freistellung',
        'text': 'Sind Sie privat krankenversichert?',
        'showIf': 'Nur sichtbar, wenn S9 = "Bezahlte Freistellung" oder "Weiterarbeiten" (nicht bei unbezahlt/keine)',
        'options': [
            {'label': 'Nein, ich bin gesetzlich versichert', 'key': 'privatVersichert', 'value': 'nein', 'color': 'green', 'next': 'S11', 'infoBox': None},
            {'label': 'Ja, ich bin privat versichert', 'key': 'privatVersichert', 'value': 'ja', 'color': 'yellow', 'next': 'S11', 'infoBox': 'Hinweis: Nach Vertragsende entfällt der Arbeitgeberzuschuss zur PKV (ca. 50% des Beitrags). Kalkulieren Sie den vollen PKV-Beitrag ab dem ersten Monat nach Beendigung ein.'},
        ],
    },
    # Kategorie 5: Zeugnis
    {
        'id': 'S11', 'category': 5, 'catLabel': 'Zeugnis',
        'text': 'Ist ein Arbeitszeugnis im Vertrag vereinbart?',
        'showIf': None,
        'options': [
            {'label': 'Ja, mit vereinbarter Note (sehr gut / gut)', 'key': 'zeugnisVereinbart', 'value': 'ja_note', 'color': 'green', 'next': 'S12', 'infoBox': None},
            {'label': 'Ja, aber ohne konkrete Notenvereinbarung', 'key': 'zeugnisVereinbart', 'value': 'ja_ohne', 'color': 'yellow', 'next': 'S12', 'infoBox': None},
            {'label': 'Nein, kein Zeugnis geregelt', 'key': 'zeugnisVereinbart', 'value': 'nein', 'color': 'red', 'next': 'S12', 'infoBox': 'Sie haben gesetzlichen Anspruch auf ein qualifiziertes Arbeitszeugnis (§109 GewO). Ohne vertragliche Regelung riskieren Sie eine schlechte Beurteilung. Bestehen Sie auf einer Vereinbarung.'},
        ],
    },
    {
        'id': 'S12', 'category': 5, 'catLabel': 'Zeugnis',
        'text': 'Ist ein konkretes Ausstellungsdatum für das Zeugnis vereinbart?',
        'showIf': None,
        'options': [
            {'label': 'Ja, konkretes Datum vereinbart', 'key': 'zeugnisdatum', 'value': 'ja', 'color': 'green', 'next': 'S13', 'infoBox': None},
            {'label': 'Nein, kein Datum vereinbart', 'key': 'zeugnisdatum', 'value': 'nein', 'color': 'yellow', 'next': 'S13', 'infoBox': 'Ohne konkretes Datum kann sich die Ausstellung des Zeugnisses verzögern — was bei Bewerbungen problematisch ist.'},
            {'label': 'Kein Zeugnis im Vertrag vereinbart', 'key': 'zeugnisdatum', 'value': 'kein_zeugnis', 'color': 'skip', 'next': 'S13', 'infoBox': None},
        ],
    },
    # Kategorie 6: Wettbewerbsverbot
    {
        'id': 'S13', 'category': 6, 'catLabel': 'Wettbewerbsverbot',
        'text': 'Enthält der Vertrag ein nachvertragliches Wettbewerbsverbot?',
        'showIf': None,
        'options': [
            {'label': 'Nein', 'key': 'wettbewerbsverbot', 'value': 'nein', 'color': 'green', 'next': 'S15', 'infoBox': None},
            {'label': 'Ja', 'key': 'wettbewerbsverbot', 'value': 'ja', 'color': 'skip', 'next': 'S14', 'infoBox': None},
            {'label': 'Ich weiß es nicht', 'key': 'wettbewerbsverbot', 'value': 'unsicher', 'color': 'yellow', 'next': 'S15', 'infoBox': None},
        ],
    },
    {
        'id': 'S14', 'category': 6, 'catLabel': 'Wettbewerbsverbot',
        'text': 'Enthält das Wettbewerbsverbot eine Karenzentschädigung?',
        'showIf': 'Nur sichtbar, wenn S13 = "Ja"',
        'options': [
            {'label': 'Ja, mindestens 50% des letzten Bruttogehalts', 'key': 'karenzentschaedigung', 'value': 'ja_50', 'color': 'green', 'next': 'S15', 'infoBox': None},
            {'label': 'Ja, aber unter 50% des letzten Bruttogehalts', 'key': 'karenzentschaedigung', 'value': 'ja_unter50', 'color': 'red', 'next': 'S15', 'infoBox': 'Ein Wettbewerbsverbot ohne Karenzentschädigung von mindestens 50% ist nach §74 HGB unverbindlich. Sie müssen es dann nicht einhalten, haben aber auch keinen Anspruch auf Entschädigung.'},
            {'label': 'Nein, keine Karenzentschädigung vereinbart', 'key': 'karenzentschaedigung', 'value': 'nein', 'color': 'red', 'next': 'S15', 'infoBox': 'Ohne Karenzentschädigung ist das Wettbewerbsverbot für Sie unverbindlich (§74 HGB). Sie können frei entscheiden, ob Sie es einhalten — haben dann aber keinen Anspruch auf Zahlung.'},
        ],
    },
    # Kategorie 7: Besonderer Schutz
    {
        'id': 'S15', 'category': 7, 'catLabel': 'Besonderer Schutz',
        'text': 'Haben Sie einen besonderen Schutzstatus?',
        'showIf': None,
        'options': [
            {'label': 'Schwerbehinderung (anerkannt oder beantragt)', 'key': 'sonderschutz', 'value': 'schwerbehindert', 'color': 'skip', 'next': 'S16', 'infoBox': None},
            {'label': 'Betriebsratsmitglied', 'key': 'sonderschutz', 'value': 'betriebsrat', 'color': 'skip', 'next': 'S16', 'infoBox': None},
            {'label': 'Schwangerschaft oder Elternzeit', 'key': 'sonderschutz', 'value': 'schwangerschaft', 'color': 'skip', 'next': 'S16', 'infoBox': None},
            {'label': 'Datenschutzbeauftragter', 'key': 'sonderschutz', 'value': 'datenschutz', 'color': 'skip', 'next': 'S16', 'infoBox': None},
            {'label': 'Nein / Keiner davon', 'key': 'sonderschutz', 'value': 'nein', 'color': 'green', 'next': 'S17', 'infoBox': None},
        ],
    },
    {
        'id': 'S16', 'category': 7, 'catLabel': 'Besonderer Schutz',
        'text': 'Wurden die erforderlichen Stellen bei Ihrem Schutzstatus einbezogen bzw. wurden Sie umfassend über Ihre Rechte informiert?',
        'showIf': 'Nur sichtbar, wenn S15 ≠ "Nein / Keiner davon"',
        'options': [
            {'label': 'Ja, ich wurde vollständig informiert und habe freiwillig zugestimmt', 'key': 'sonderschutzEinbezogen', 'value': 'ja', 'color': 'green', 'next': 'S17', 'infoBox': None},
            {'label': 'Nein / Ich bin nicht sicher', 'key': 'sonderschutzEinbezogen', 'value': 'nein', 'color': 'red', 'next': 'S17', 'infoBox': 'Besonders wichtig: Bei Schwerbehinderten braucht der Arbeitgeber für eine Kündigung die Zustimmung des Integrationsamts (§168 SGB IX). Beim Aufhebungsvertrag können Sie freiwillig zustimmen — aber nur nach vollständiger Information über Ihre Rechte. Lassen Sie dies unbedingt prüfen.'},
        ],
    },
    # Kategorie 8: Verhandlung & Druck
    {
        'id': 'S17', 'category': 8, 'catLabel': 'Verhandlung & Druck',
        'text': 'Wie lange hatten Sie Zeit, den Vertrag zu prüfen, bevor Sie unterschreiben sollten?',
        'showIf': None,
        'options': [
            {'label': 'Mehrere Tage oder länger', 'key': 'bedenkzeit', 'value': 'tage', 'color': 'green', 'next': 'S18', 'infoBox': None},
            {'label': '1–2 Tage', 'key': 'bedenkzeit', 'value': '1_2', 'color': 'yellow', 'next': 'S18', 'infoBox': None},
            {'label': 'Ich sollte sofort unterschreiben', 'key': 'bedenkzeit', 'value': 'sofort', 'color': 'red', 'next': 'S18', 'infoBox': 'Das BAG hat in BAG 6 AZR 333/21 klargestellt: Sofortiger Unterschriftsdruck allein ist noch kein Verstoß gegen das Gebot fairen Verhandelns — aber in Kombination mit anderen Druckmitteln kann der Vertrag anfechtbar sein.'},
        ],
    },
    {
        'id': 'S18', 'category': 8, 'catLabel': 'Verhandlung & Druck',
        'text': 'Wurde Ihnen eine Kündigung oder andere Nachteile angedroht für den Fall, dass Sie nicht unterschreiben?',
        'showIf': None,
        'options': [
            {'label': 'Nein, keine Drohungen', 'key': 'druckAusgeubt', 'value': 'nein', 'color': 'green', 'next': 'ERGEBNIS', 'infoBox': None},
            {'label': 'Ja, mit Kündigung gedroht', 'key': 'druckAusgeubt', 'value': 'kuendigung', 'color': 'yellow', 'next': 'ERGEBNIS', 'infoBox': 'Eine Drohung mit Kündigung ist nicht automatisch unzulässig — nur wenn der Arbeitgeber keine realistische Grundlage für eine Kündigung hatte. Lassen Sie dies prüfen.'},
            {'label': 'Ja, mit anderen Nachteilen gedroht', 'key': 'druckAusgeubt', 'value': 'andere', 'color': 'red', 'next': 'ERGEBNIS', 'infoBox': 'Widerrechtliche Drohung berechtigt zur Anfechtung des Aufhebungsvertrags (§123 BGB). Handeln Sie schnell — die Anfechtungsfrist beträgt 1 Jahr.'},
        ],
    },
]

current_cat = None
for q in questions:
    if q['category'] != current_cat:
        current_cat = q['category']
        doc.add_heading(f"Kategorie {current_cat} — {q['catLabel']}", level=2)

    doc.add_heading(f"{q['id']}: {q['text']}", level=3)

    # Help text
    if q.get('help'):
        p = doc.add_paragraph()
        run = p.add_run('Erklärungstext: ')
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run = p.add_run(q['help'])
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True

    # Conditional logic
    if q.get('showIf'):
        p = doc.add_paragraph()
        run = p.add_run('\u26A0 Bedingte Anzeige: ')
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
        run = p.add_run(q['showIf'])
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

    # Options table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, text in enumerate(['Antworttext', 'Ampel', 'Nächster Schritt', 'InfoBox?', 'Datenwert']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
        set_cell_bg(hdr.cells[i], 'F0EAD9')

    for opt in q['options']:
        row = table.add_row()

        row.cells[0].text = opt['label']
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)

        row.cells[1].text = ''
        p = row.cells[1].paragraphs[0]
        run = p.add_run(color_labels[opt['color']])
        run.font.size = Pt(8.5)
        run.bold = True
        set_cell_bg(row.cells[1], color_map[opt['color']])

        row.cells[2].text = f"→ {opt['next']}"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(8.5)

        has_info = 'Ja' if opt['infoBox'] else '—'
        row.cells[3].text = has_info
        row.cells[3].paragraphs[0].runs[0].font.size = Pt(8.5)

        row.cells[4].text = f"{opt['key']} = {opt['value']}"
        row.cells[4].paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph('')

    # InfoBox contents
    info_opts = [opt for opt in q['options'] if opt.get('infoBox')]
    if info_opts:
        for opt in info_opts:
            p = doc.add_paragraph()
            run = p.add_run(f'InfoBox bei "{opt["label"]}": ')
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7A, 0x65, 0x28)
            run = p.add_run(opt['infoBox'])
            run.font.size = Pt(9)
            run.italic = True

    # Comment area
    p = doc.add_paragraph()
    run = p.add_run(f'Kommentar zu {q["id"]}: ')
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run('_______________________________________________')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 3. BEWERTUNGSLOGIK
# ══════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading('3. Bewertungslogik — Gesamtbewertung', level=1)

doc.add_paragraph(
    'Die Bewertung erfolgt in zwei Ebenen: pro Kategorie und als Gesamtwert.'
)

doc.add_heading('Punktevergabe', level=2)
doc.add_paragraph(
    '  \u2022 Grün = 2 Punkte\n'
    '  \u2022 Gelb = 1 Punkt\n'
    '  \u2022 Rot = 0 Punkte\n'
    '  \u2022 Skip = wird nicht in die Berechnung einbezogen'
)

doc.add_heading('Gesamtprozent', level=2)
doc.add_paragraph(
    'Gesamtprozent = (Summe aller Punkte) / (Anzahl bewerteter Fragen × 2) × 100\n\n'
    'Fragen mit "Skip" werden aus der Berechnung ausgeschlossen.'
)

doc.add_heading('Kategorie-Bewertung', level=2)
doc.add_paragraph(
    'Für jede Kategorie wird der Durchschnitt der Punkte berechnet:\n'
    '  \u2022 Durchschnitt ≥ 0,75 → Grün\n'
    '  \u2022 Durchschnitt ≥ 0,40 → Gelb\n'
    '  \u2022 Durchschnitt < 0,40 → Rot\n\n'
    'Kategorien ohne bewertete Fragen werden als Grün dargestellt.'
)

doc.add_heading('Gesamtverdikt', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0]
for i, text in enumerate(['Gesamtprozent', 'Ampelfarbe', 'Titel']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

verdicts = [
    ('≥ 80%', 'GRÜN', 'Der Vertrag erscheint überwiegend fair.'),
    ('50–79%', 'GELB', 'Der Vertrag enthält Verbesserungsbedarf.'),
    ('< 50%', 'ROT', 'Der Vertrag enthält kritische Punkte.'),
]

verdict_bg = {'GRÜN': 'D4EDDA', 'GELB': 'FFF3CC', 'ROT': 'FFCCCC'}

for v in verdicts:
    row = table.add_row()
    for i, text in enumerate(v):
        row.cells[i].text = text
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(row.cells[1], verdict_bg[v[1]])

doc.add_paragraph('')

doc.add_heading('Ergebnis-Texte je Verdikt', level=2)

result_texts = [
    ('GRÜN', 'Die von Ihnen gemachten Angaben deuten auf faire Vertragsbedingungen hin. Dennoch empfehlen wir eine kurze anwaltliche Prüfung vor Unterzeichnung.', 'Gut — aber sicher ist sicher.', 'Auch bei einem fairen Vertrag lohnt sich eine kurze anwaltliche Durchsicht. Nutzen Sie die kostenlose Erstberatung durch Fachanwalt Fatih Bektas.'),
    ('GELB', 'Mehrere Punkte sind verbesserungswürdig. Eine Nachverhandlung erscheint sinnvoll — wir zeigen Ihnen, was möglich ist.', 'Nachverhandeln ist möglich.', 'Mehrere Punkte bieten Verbesserungspotenzial. Fachanwalt Fatih Bektas zeigt Ihnen in einer kostenlosen Erstberatung, was Sie nachverhandeln können — Antwort in 24 Stunden.'),
    ('ROT', 'Mehrere Klauseln sind problematisch. Wir empfehlen dringend eine anwaltliche Prüfung vor Unterzeichnung.', 'Bitte nicht ohne Beratung unterschreiben.', 'Ihr Vertrag enthält kritische Punkte. Bevor Sie unterschreiben, nutzen Sie die kostenlose Erstberatung durch Fachanwalt Fatih Bektas — Antwort innerhalb von 24 Stunden.'),
]

for v_color, v_text, v_cta_headline, v_cta_text in result_texts:
    p = doc.add_paragraph()
    run = p.add_run(f'{v_color}: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(v_text)
    run.font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'CTA-Headline: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(v_cta_headline)
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'CTA-Text: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(v_cta_text)
    run.font.size = Pt(9)

    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# 4. BEDINGTE LOGIK
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('4. Bedingte Logik & Navigation', level=1)

doc.add_paragraph(
    'Der Checker zeigt Fragen in einer festen Reihenfolge, aber mit bedingten Verzweigungen:'
)

cond_table = doc.add_table(rows=1, cols=3)
cond_table.style = 'Table Grid'
hdr = cond_table.rows[0]
for i, text in enumerate(['Frage', 'Wird angezeigt wenn...', 'Sonst → nächster Schritt']):
    hdr.cells[i].text = text
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_bg(hdr.cells[i], 'F0EAD9')

conditionals = [
    ('S2', 'S1 = "Ja, eine Abfindung ist vereinbart"', '→ S19 (Turboklausel)'),
    ('S3', 'S1 = "Ja" (über S2 erreichbar)', '→ S19'),
    ('S10', 'S9 = "Bezahlte Freistellung" oder "Weiterarbeiten"', '→ S11 (bei unbezahlt/keine)'),
    ('S14', 'S13 = "Ja" (Wettbewerbsverbot vorhanden)', '→ S15 (bei Nein/Unsicher)'),
    ('S16', 'S15 ≠ "Nein / Keiner davon"', '→ S17 (bei kein Sonderschutz)'),
]

for c in conditionals:
    add_table_row(cond_table, c, bold_first=True)

doc.add_paragraph('')

doc.add_heading('Navigationsfluss', level=2)
doc.add_paragraph(
    'Standard-Pfad (alle Fragen sichtbar):\n'
    'S1 → S2 → S3 → S19 → S4 → S5 → S6 → S7 → S8 → S9 → S10 → S11 → S12 → S13 → S14 → S15 → S16 → S17 → S18 → ERGEBNIS\n\n'
    'Kürzester Pfad (keine Abfindung, kein Wettbewerb, kein Sonderschutz):\n'
    'S1 → S19 → S4 → S5 → S6 → S7 → S8 → S9 → S11 → S12 → S13 → S15 → S17 → S18 → ERGEBNIS\n'
    '(14 Fragen statt 19)'
)

# ══════════════════════════════════════════════════════════════════════
# 5. FAQ
# ══════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading('5. FAQ-Inhalte (auf der Seite angezeigt + Schema.org)', level=1)

faqs = [
    ('Muss ich einen Aufhebungsvertrag unterschreiben?', 'Nein. Ein Aufhebungsvertrag kommt nur durch beidseitige Zustimmung zustande. Sie können ablehnen — dann muss der Arbeitgeber kündigen, wenn er das Arbeitsverhältnis beenden will. Eine Kündigung gibt Ihnen mehr Rechte: 3-Wochen-Klagefrist, KSchG-Schutz und kein automatisches Sperrzeit-Risiko beim ALG I.'),
    ('Kann ich einen Aufhebungsvertrag rückgängig machen?', 'Nur unter engen Voraussetzungen: bei widerrechtlicher Drohung (§123 BGB, Frist: 1 Jahr) oder bei Verstoß gegen das Gebot fairen Verhandelns (BAG 6 AZR 75/18). Ein gesetzliches Widerrufsrecht wie bei Verbraucherverträgen gibt es nicht — es sei denn, es ist vertraglich vereinbart.'),
    ('Droht immer eine Sperrzeit beim Arbeitslosengeld?', 'Grundsätzlich ja. Bei Aufhebungsverträgen verhängt die Agentur für Arbeit regelmäßig eine 12-wöchige Sperrzeit (§159 SGB III), weil der Arbeitnehmer das Arbeitsverhältnis „selbst aufgelöst" hat. Ausnahmen: drohende betriebsbedingte Kündigung oder wichtiger persönlicher Grund.'),
    ('Was ist das Gebot fairen Verhandelns?', 'Das BAG hat entschieden, dass Arbeitgeber beim Abschluss von Aufhebungsverträgen fair verhandeln müssen. Eine psychische Drucksituation kann zur Unwirksamkeit führen (BAG 6 AZR 75/18, BAG 6 AZR 333/21).'),
    ('Wie hoch sollte die Abfindung sein?', 'Es gibt keinen gesetzlichen Anspruch auf eine bestimmte Abfindungshöhe. Faustformel: 0,5 Bruttomonatsgehälter pro Beschäftigungsjahr. Bei starkem Kündigungsschutz oder Verfahrensfehlern sind deutlich höhere Beträge erzielbar.'),
    ('Was gilt beim Wettbewerbsverbot ohne Karenzentschädigung?', 'Ein nachvertragliches Wettbewerbsverbot ohne Karenzentschädigung von mindestens 50% des letzten Bruttogehalts ist nach §74 HGB unverbindlich.'),
]

for i, (fq, fa) in enumerate(faqs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'FAQ {i}: {fq}')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph(fa)
    p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.size = Pt(9.5)

    p = doc.add_paragraph()
    run = p.add_run(f'Kommentar: ')
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run('_______________________________________________')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Ende des Entscheidungsbaums —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('APOS Legal  ·  gekuendigt-abfindung.de  ·  Generiert am: April 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_dir = os.path.expanduser('~/Desktop')
output_path = os.path.join(output_dir, 'Aufhebungsvertrag-Checker-Entscheidungsbaum.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
